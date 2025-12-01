import os
import io
import json
import traceback
import threading
import queue
import uuid
import time
import math
from datetime import datetime
from faster_whisper import WhisperModel
from llama_cpp import Llama
from docx import Document
from extensions import db
from models import Transcript, User, TranscriptionTask

# Global model instances
whisper_model = None
llm_model = None
MODEL_FILENAME = "Llama-3.2-3B-Instruct-Q4_K_M.gguf"

# Task Queue System
# We still use a memory queue for the worker to pick up jobs, 
# but the STATE is stored in the DB.
task_queue = queue.Queue()

class BackgroundWorker(threading.Thread):
    def __init__(self, app):
        super().__init__()
        self.daemon = True
        self.running = True
        self.app = app

    def run(self):
        print("Background Worker Started")
        load_models()
        
        while self.running:
            try:
                task_data = task_queue.get(timeout=1)
                task_id = task_data['id']
                audio_path = task_data['audio_path']
                
                with self.app.app_context():
                    try:
                        # Update status to processing
                        update_task_status(task_id, "processing", 0, "Memulai proses...")
                        
                        # 1. Transcribe
                        update_task_status(task_id, "processing", 10, "Mentranskripsi audio...")
                        raw_transcript = transcribe_audio(audio_path)
                        
                        # 2. Format Dialogue (Chunked)
                        update_task_status(task_id, "processing", 40, "Memformat dialog...")
                        formatted_content = format_dialogue_chunked(raw_transcript, task_id)
                        
                        # 3. Extract Metadata
                        update_task_status(task_id, "processing", 80, "Mengekstrak informasi...")
                        metadata = extract_metadata_from_transcript(formatted_content)
                        
                        # 4. Save to DB (Transcript)
                        task = TranscriptionTask.query.get(task_id)
                        if task:
                            new_transcript = Transcript(
                                user_id=task.user_id,
                                filename=os.path.basename(audio_path),
                                participant_code=metadata.get('participant_code', '-'),
                                participant_name=metadata.get('participant_name', '-'),
                                participant_age=metadata.get('participant_age', '-'),
                                participant_education=metadata.get('participant_education', '-'),
                                content=formatted_content
                            )
                            db.session.add(new_transcript)
                            db.session.commit()
                            
                            # Update Task to Completed
                            task.status = 'completed'
                            task.progress = 100
                            task.message = 'Selesai'
                            task.result_id = new_transcript.id
                            db.session.commit()
                            
                            # Clean up file
                            if os.path.exists(audio_path):
                                os.remove(audio_path)
                        
                    except Exception as e:
                        print(f"Task failed: {e}")
                        traceback.print_exc()
                        update_task_status(task_id, "failed", 0, str(e), error=str(e))
                
                task_queue.task_done()
                
            except queue.Empty:
                continue
            except Exception as e:
                print(f"Worker error: {e}")

def start_worker(app):
    worker = BackgroundWorker(app)
    worker.start()
    return worker

def load_models():
    global whisper_model, llm_model
    
    if whisper_model is None:
        print("Loading Whisper model...")
        # Use 'medium' model as requested
        model_path = "models/whisper-medium" 
        if os.path.exists(model_path):
             whisper_model = WhisperModel(model_path, device="cpu", compute_type="int8")
        else:
             # Fallback or auto-download if setup_models.py wasn't run
             whisper_model = WhisperModel("medium", device="cpu", compute_type="int8")
            
    if llm_model is None:
        print("Loading LLM model...")
        model_path = f"models/{MODEL_FILENAME}"
        if os.path.exists(model_path):
            llm_model = Llama(
                model_path=model_path,
                n_ctx=4096,
                n_threads=4
            )
        else:
            print("LLM Model not found locally.")

def add_task(audio_path, user_id):
    task_id = str(uuid.uuid4())
    
    # Create DB Entry
    new_task = TranscriptionTask(
        id=task_id,
        user_id=user_id,
        filename=os.path.basename(audio_path),
        status='queued',
        progress=0,
        message='Menunggu antrian...'
    )
    db.session.add(new_task)
    db.session.commit()
    
    # Add to memory queue for worker
    task_queue.put({
        'id': task_id,
        'audio_path': audio_path,
        'user_id': user_id
    })
    
    return task_id

def update_task_status(task_id, status, progress, message, error=None):
    # This function must be called within an app context
    try:
        task = TranscriptionTask.query.get(task_id)
        if task:
            task.status = status
            task.progress = progress
            task.message = message
            if error:
                task.error = error
            db.session.commit()
    except Exception as e:
        print(f"Failed to update task status: {e}")

def get_task_status(task_id):
    task = TranscriptionTask.query.get(task_id)
    if not task:
        return None
    
    result_data = None
    if task.status == 'completed' and task.result_id:
        transcript = Transcript.query.get(task.result_id)
        if transcript:
            result_data = {
                "transcript_id": transcript.id,
                "content": transcript.content,
                "metadata": {
                    "participant_code": transcript.participant_code,
                    "participant_name": transcript.participant_name,
                    "participant_age": transcript.participant_age,
                    "participant_education": transcript.participant_education
                }
            }
            
    return {
        'id': task.id,
        'status': task.status,
        'progress': task.progress,
        'message': task.message,
        'error': task.error,
        'result': result_data
    }

def transcribe_audio(audio_path):
    segments, info = whisper_model.transcribe(audio_path, beam_size=5)
    
    full_text = ""
    for segment in segments:
        full_text += segment.text + " "
        
    return full_text.strip()

def format_dialogue_chunked(text, task_id):
    if not llm_model:
        return text
        
    # Simple chunking logic
    chunk_size = 2000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    
    formatted_chunks = []
    total_chunks = len(chunks)
    
    for i, chunk in enumerate(chunks):
        # Update progress based on chunk processing
        progress = 40 + int((i / total_chunks) * 40) # 40% to 80%
        update_task_status(task_id, "processing", progress, f"Memformat bagian {i+1}/{total_chunks}...")
        
        prompt = f"""
        Ubah teks berikut menjadi format dialog wawancara yang rapi.
        Tandai pembicara dengan "Q:" (Pewawancara) dan "A:" (Partisipan) jika bisa dideteksi.
        Jika tidak, rapikan saja tanda bacanya.
        
        Teks:
        {chunk}
        
        Dialog:
        """
        
        output = llm_model(
            prompt, 
            max_tokens=1024, 
            stop=["Teks:", "Dialog:"], 
            echo=False
        )
        formatted_chunks.append(output['choices'][0]['text'].strip())
        
    return "\n\n".join(formatted_chunks)

def extract_metadata_from_transcript(text):
    if not llm_model:
        return {}
        
    prompt = f"""
    Analisis transkrip berikut dan ekstrak informasi:
    1. Kode Partisipan (contoh: P1, P2)
    2. Nama Partisipan
    3. Usia
    4. Pendidikan Terakhir
    
    Jika tidak ditemukan, tulis "-".
    
    Format output JSON:
    {{
        "participant_code": "...",
        "participant_name": "...",
        "participant_age": "...",
        "participant_education": "..."
    }}
    
    Transkrip (awal):
    {text[:4000]}
    """
    
    try:
        output = llm_model(
            prompt,
            max_tokens=200,
            stop=["Transkrip:"],
            echo=False
        )
        json_str = output['choices'][0]['text'].strip()
        # Clean up JSON string if needed
        if "```json" in json_str:
            json_str = json_str.split("```json")[1].split("```")[0]
        elif "{" not in json_str:
             # Fallback if LLM doesn't output JSON
             return {}
             
        return json.loads(json_str)
    except:
        return {}

def generate_docx(transcript):
    doc = Document()
    doc.add_heading('TRANSKRIP WAWANCARA', 0)
    
    p = doc.add_paragraph()
    p.add_run('Kode Partisipan: ').bold = True
    p.add_run(f"{transcript.participant_code}\n")
    p.add_run('Nama: ').bold = True
    p.add_run(f"{transcript.participant_name}\n")
    p.add_run('Usia: ').bold = True
    p.add_run(f"{transcript.participant_age}\n")
    p.add_run('Pendidikan: ').bold = True
    p.add_run(f"{transcript.participant_education}\n")
    p.add_run('Waktu: ').bold = True
    p.add_run(f"{transcript.created_at}\n")
    
    doc.add_heading('Isi Transkrip', level=1)
    doc.add_paragraph(transcript.content)
    
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem
